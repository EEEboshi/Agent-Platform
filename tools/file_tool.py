"""
文件处理工具
支持读取/写入 PDF、Excel、Word 文件
"""
from llama_index.core.tools import FunctionTool
from pydantic import BaseModel, Field
from typing import Optional, List
import logging
import os
from pathlib import Path

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


class FileReadInput(BaseModel):
    """文件读取输入验证模型"""
    file_path: str = Field(description="要读取的文件路径")
    sheet_name: Optional[str] = Field(default=None, description="Excel 工作表名称（仅 Excel 文件使用）")


class FileWriteInput(BaseModel):
    """文件写入输入验证模型"""
    file_path: str = Field(description="要写入的文件路径")
    content: str = Field(description="要写入的内容")
    file_type: str = Field(description="文件类型：txt, csv, xlsx, docx, pdf")


class ExcelReadInput(BaseModel):
    """Excel 读取输入验证模型"""
    file_path: str = Field(description="Excel 文件路径")
    sheet_name: Optional[str] = Field(default=0, description="工作表名称或索引")
    header_row: int = Field(default=0, description="表头行号")


class ExcelWriteInput(BaseModel):
    """Excel 写入输入验证模型"""
    file_path: str = Field(description="Excel 文件路径")
    data: str = Field(description="数据（JSON 格式或 CSV 格式）")
    sheet_name: str = Field(default="Sheet1", description="工作表名称")


class FileValidator:
    """文件验证器"""
    
    ALLOWED_EXTENSIONS = {
        'read': ['.txt', '.csv', '.xlsx', '.xls', '.docx', '.doc', '.pdf', '.json', '.md'],
        'write': ['.txt', '.csv', '.xlsx', '.docx', '.pdf', '.json', '.md']
    }
    
    @staticmethod
    def validate_file_path(file_path: str, operation: str = 'read') -> tuple[bool, str]:
        """验证文件路径"""
        if not file_path:
            return False, "文件路径不能为空"
        
        path = Path(file_path)
        
        # 检查扩展名
        ext = path.suffix.lower()
        if operation == 'read' and ext not in FileValidator.ALLOWED_EXTENSIONS['read']:
            return False, f"不支持读取的文件类型：{ext}"
        if operation == 'write' and ext not in FileValidator.ALLOWED_EXTENSIONS['write']:
            return False, f"不支持写入的文件类型：{ext}"
        
        # 读取操作检查文件是否存在
        if operation == 'read' and not path.exists():
            return False, f"文件不存在：{file_path}"
        
        # 写入操作检查目录是否存在
        if operation == 'write':
            parent_dir = path.parent
            if not parent_dir.exists():
                try:
                    parent_dir.mkdir(parents=True, exist_ok=True)
                    logger.info(f"创建目录：{parent_dir}")
                except Exception as e:
                    return False, f"无法创建目录：{e}"
        
        return True, ""


def read_text_file(file_path: str) -> str:
    """
    读取文本文件内容。
    支持 .txt, .csv, .md, .json 等文本格式
    
    Args:
        file_path: 文件路径
        
    Returns:
        文件内容
    """
    is_valid, error_msg = FileValidator.validate_file_path(file_path, 'read')
    if not is_valid:
        return f"错误：{error_msg}"
    
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()
        logger.info(f"成功读取文件：{file_path}")
        return content
    except Exception as e:
        logger.error(f"读取文件失败：{e}")
        return f"读取失败：{str(e)}"


def write_text_file(file_path: str, content: str) -> str:
    """
    写入文本文件。
    支持 .txt, .csv, .md, .json 等文本格式
    
    Args:
        file_path: 文件路径
        content: 要写入的内容
        
    Returns:
        操作结果
    """
    is_valid, error_msg = FileValidator.validate_file_path(file_path, 'write')
    if not is_valid:
        return f"错误：{error_msg}"
    
    try:
        with open(file_path, 'w', encoding='utf-8') as f:
            f.write(content)
        logger.info(f"成功写入文件：{file_path}")
        return f"文件已成功写入：{file_path}"
    except Exception as e:
        logger.error(f"写入文件失败：{e}")
        return f"写入失败：{str(e)}"


def read_excel(file_path: str, sheet_name: Optional[str] = None) -> str:
    """
    读取 Excel 文件内容。
    支持 .xlsx, .xls 格式
    
    Args:
        file_path: Excel 文件路径
        sheet_name: 工作表名称或索引（默认第一个工作表）
        
    Returns:
        Excel 数据（CSV 格式）
    """
    try:
        import pandas as pd
        
        is_valid, error_msg = FileValidator.validate_file_path(file_path, 'read')
        if not is_valid:
            return f"错误：{error_msg}"
        
        # 读取 Excel
        if sheet_name is None:
            df = pd.read_excel(file_path)
        else:
            try:
                sheet_idx = int(sheet_name)
                df = pd.read_excel(file_path, sheet_name=sheet_idx)
            except ValueError:
                df = pd.read_excel(file_path, sheet_name=sheet_name)
        
        # 转换为 CSV 格式字符串
        csv_data = df.to_csv(index=False)
        logger.info(f"成功读取 Excel 文件：{file_path}, 行数：{len(df)}")
        return csv_data
        
    except ImportError:
        return "错误：需要安装 pandas 和 openpyxl，请运行：pip install pandas openpyxl"
    except Exception as e:
        logger.error(f"读取 Excel 失败：{e}")
        return f"读取失败：{str(e)}"


def write_excel(file_path: str, data: str, sheet_name: str = "Sheet1") -> str:
    """
    写入 Excel 文件。
    支持将 CSV 格式或 JSON 格式数据写入 Excel
    
    Args:
        file_path: Excel 文件路径
        data: 数据（CSV 格式或 JSON 格式）
        sheet_name: 工作表名称
        
    Returns:
        操作结果
    """
    try:
        import pandas as pd
        import json
        
        is_valid, error_msg = FileValidator.validate_file_path(file_path, 'write')
        if not is_valid:
            return f"错误：{error_msg}"
        
        # 尝试解析数据
        try:
            # 先尝试 JSON 格式
            data_dict = json.loads(data)
            if isinstance(data_dict, list):
                df = pd.DataFrame(data_dict)
            else:
                df = pd.DataFrame([data_dict])
        except json.JSONDecodeError:
            # 尝试 CSV 格式
            from io import StringIO
            df = pd.read_csv(StringIO(data))
        
        # 写入 Excel
        with pd.ExcelWriter(file_path, engine='openpyxl') as writer:
            df.to_excel(writer, sheet_name=sheet_name, index=False)
        
        logger.info(f"成功写入 Excel 文件：{file_path}, 行数：{len(df)}")
        return f"Excel 文件已成功写入：{file_path}, 共 {len(df)} 行数据"
        
    except ImportError:
        return "错误：需要安装 pandas 和 openpyxl，请运行：pip install pandas openpyxl"
    except Exception as e:
        logger.error(f"写入 Excel 失败：{e}")
        return f"写入失败：{str(e)}"


def read_word(file_path: str) -> str:
    """
    读取 Word 文档内容。
    支持 .docx 格式
    
    Args:
        file_path: Word 文件路径
        
    Returns:
        Word 文档文本内容
    """
    try:
        from docx import Document
        
        is_valid, error_msg = FileValidator.validate_file_path(file_path, 'read')
        if not is_valid:
            return f"错误：{error_msg}"
        
        doc = Document(file_path)
        text_content = []
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_content.append(paragraph.text)
        
        content = '\n'.join(text_content)
        logger.info(f"成功读取 Word 文档：{file_path}, 段落数：{len(doc.paragraphs)}")
        return content
        
    except ImportError:
        return "错误：需要安装 python-docx，请运行：pip install python-docx"
    except Exception as e:
        logger.error(f"读取 Word 文档失败：{e}")
        return f"读取失败：{str(e)}"


def write_word(file_path: str, content: str) -> str:
    """
    写入 Word 文档。
    支持创建 .docx 格式文档
    
    Args:
        file_path: Word 文件路径
        content: 文档内容（段落用换行符分隔）
        
    Returns:
        操作结果
    """
    try:
        from docx import Document
        
        is_valid, error_msg = FileValidator.validate_file_path(file_path, 'write')
        if not is_valid:
            return f"错误：{error_msg}"
        
        doc = Document()
        
        # 按段落分割内容
        paragraphs = content.split('\n')
        for para in paragraphs:
            if para.strip():
                doc.add_paragraph(para)
        
        doc.save(file_path)
        logger.info(f"成功写入 Word 文档：{file_path}, 段落数：{len(paragraphs)}")
        return f"Word 文档已成功写入：{file_path}"
        
    except ImportError:
        return "错误：需要安装 python-docx，请运行：pip install python-docx"
    except Exception as e:
        logger.error(f"写入 Word 文档失败：{e}")
        return f"写入失败：{str(e)}"


def read_pdf(file_path: str) -> str:
    """
    读取 PDF 文件内容。
    支持 .pdf 格式
    
    Args:
        file_path: PDF 文件路径
        
    Returns:
        PDF 文本内容
    """
    try:
        import PyPDF2
        
        is_valid, error_msg = FileValidator.validate_file_path(file_path, 'read')
        if not is_valid:
            return f"错误：{error_msg}"
        
        text_content = []
        
        with open(file_path, 'rb') as f:
            reader = PyPDF2.PdfReader(f)
            logger.info(f"PDF 页数：{len(reader.pages)}")
            
            for i, page in enumerate(reader.pages):
                text = page.extract_text()
                if text.strip():
                    text_content.append(f"--- 第 {i+1} 页 ---")
                    text_content.append(text)
        
        content = '\n'.join(text_content)
        logger.info(f"成功读取 PDF 文件：{file_path}")
        return content
        
    except ImportError:
        return "错误：需要安装 PyPDF2，请运行：pip install PyPDF2"
    except Exception as e:
        logger.error(f"读取 PDF 失败：{e}")
        return f"读取失败：{str(e)}"


# 创建工具实例
read_text_tool = FunctionTool.from_defaults(
    fn=read_text_file,
    name="read_text_file",
    description="读取文本文件内容（.txt, .csv, .md, .json 等）",
    fn_schema=FileReadInput
)

write_text_tool = FunctionTool.from_defaults(
    fn=write_text_file,
    name="write_text_file",
    description="写入文本文件（.txt, .csv, .md, .json 等）",
    fn_schema=FileWriteInput
)

read_excel_tool = FunctionTool.from_defaults(
    fn=read_excel,
    name="read_excel",
    description="读取 Excel 文件内容（.xlsx, .xls）",
    fn_schema=ExcelReadInput
)

write_excel_tool = FunctionTool.from_defaults(
    fn=write_excel,
    name="write_excel",
    description="写入 Excel 文件，支持 CSV 或 JSON 格式数据",
    fn_schema=ExcelWriteInput
)

read_word_tool = FunctionTool.from_defaults(
    fn=read_word,
    name="read_word",
    description="读取 Word 文档内容（.docx）",
    fn_schema=FileReadInput
)

write_word_tool = FunctionTool.from_defaults(
    fn=write_word,
    name="write_word",
    description="写入 Word 文档（.docx）",
    fn_schema=FileWriteInput
)

read_pdf_tool = FunctionTool.from_defaults(
    fn=read_pdf,
    name="read_pdf",
    description="读取 PDF 文件内容",
    fn_schema=FileReadInput
)

# 汇总所有文件处理工具
FILE_TOOLS = [
    read_text_tool,
    write_text_tool,
    read_excel_tool,
    write_excel_tool,
    read_word_tool,
    write_word_tool,
    read_pdf_tool
]
